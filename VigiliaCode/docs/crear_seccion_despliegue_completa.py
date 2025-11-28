# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ============================================================
# SECCION: DIAGRAMA DE DESPLIEGUE (ACTUALIZADA)
# ============================================================

doc.add_heading('Diagrama de despliegue', level=1)

doc.add_paragraph(
    'Este diagrama muestra como se distribuye fisicamente el sistema en su infraestructura, '
    'indicando donde se ejecuta cada componente. Facilita la comunicacion entre los distintos '
    'equipos de desarrollo, ademas de analizar el rendimiento, la disponibilidad y la escalabilidad del sistema.'
)

doc.add_paragraph(
    'En el diagrama se muestran los principales componentes que permiten la ejecucion optima e integrada '
    'de todo el sistema VigilIA, actualmente desplegado y funcionando en produccion:'
)

# ============================================================
# 1. GOOGLE CLOUD PLATFORM
# ============================================================
doc.add_heading('1. Google Cloud Platform (GCP)', level=2)

doc.add_paragraph(
    'La plataforma en la cual se aloja toda la infraestructura del sistema. Permite un despliegue seguro, '
    'escalable y con servicios integrados que simplifican la operacion y el mantenimiento. GCP provee el '
    'ecosistema donde se ejecutan los servicios, bases de datos, almacenamiento y mensajeria.'
)

doc.add_paragraph('Configuracion del proyecto:')
bullets = [
    'Proyecto GCP: composed-apogee-475623-p6',
    'Region principal: southamerica-west1 (Santiago, Chile)',
    'Zona de disponibilidad: southamerica-west1-a'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 2. FIREBASE SERVICES
# ============================================================
doc.add_heading('2. Firebase Services', level=2)

doc.add_paragraph(
    'Firebase proporciona servicios esenciales para la autenticacion de usuarios y el alojamiento '
    'de la aplicacion web. Esta completamente integrado con GCP y permite una experiencia de usuario fluida.'
)

doc.add_heading('2.1 Firebase Hosting', level=3)
doc.add_paragraph('Servicio de alojamiento web donde se despliega el frontend de la aplicacion. Estado actual: OPERATIVO')
bullets = [
    'URL principal: https://app-vigilia.web.app',
    'Dominio personalizado: https://app.mivigilia.cl',
    'HTTPS automatico con certificado SSL gratuito',
    'CDN global para baja latencia en toda Latinoamerica',
    'Despliegue continuo integrado con el repositorio del proyecto'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2.2 Firebase Authentication', level=3)
doc.add_paragraph('Sistema de autenticacion de usuarios completamente funcional. Estado actual: OPERATIVO')
bullets = [
    'Registro e inicio de sesion con email y contrasena',
    'Generacion y validacion de tokens JWT para cada sesion',
    'Integracion nativa con el frontend React Native/Expo',
    'Persistencia de sesion segura en dispositivos moviles',
    'Soporte para recuperacion de contrasena via email'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 3. CLOUD RUN - MICROSERVICIOS
# ============================================================
doc.add_heading('3. Cloud Run (Arquitectura Serverless)', level=2)

doc.add_paragraph(
    'Plataforma serverless que ejecuta los microservicios del backend en contenedores Docker con '
    'escalamiento automatico. Esta arquitectura permite que el sistema escale de 0 a N instancias '
    'segun la demanda, optimizando costos y garantizando disponibilidad.'
)

doc.add_paragraph('Caracteristicas principales de Cloud Run:')
bullets = [
    'Escalamiento automatico de 0 a N instancias segun trafico',
    'Pago solo por uso (sin costo cuando no hay trafico)',
    'Aislamiento seguro mediante contenedores Docker',
    'Despliegue continuo desde Cloud Build',
    'HTTPS automatico con certificados gestionados por Google',
    'Tiempo de arranque en frio menor a 5 segundos'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('El sistema cuenta con 4 microservicios desplegados y operativos:')

# 3.1 api-backend
doc.add_heading('3.1 api-backend (FastAPI) - Servicio Principal', level=3)
doc.add_paragraph('Estado actual: OPERATIVO - Revision: api-backend-00113')
doc.add_paragraph(
    'Componente central del procesamiento logico del sistema. Es el nucleo que coordina todas las '
    'operaciones entre el frontend, la base de datos y los servicios externos.'
)
doc.add_paragraph('URL: https://api-backend-687053793381.southamerica-west1.run.app')
doc.add_paragraph('Funcionalidades implementadas y funcionando:')
bullets = [
    'Endpoints REST para comunicacion con el frontend (mas de 30 endpoints)',
    'Gestion completa de usuarios (registro, login, perfiles, roles)',
    'Administracion de adultos mayores y relaciones con cuidadores',
    'Sistema de alertas: recepcion, almacenamiento y distribucion de alertas de caidas',
    'Sistema de solicitudes de ayuda manual desde la aplicacion',
    'Gestion de recordatorios (medicamentos, citas medicas, ejercicio)',
    'Control de dispositivos NanoPi registrados',
    'Configuracion de preferencias de notificacion por usuario',
    'Integracion con Cloud Storage para imagenes de evidencia',
    'Comunicacion segura con Secret Manager para credenciales',
    'Validacion de tokens Firebase para autenticacion',
    'Sistema de suscripciones y planes de pago'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.2 alertas-websocket
doc.add_heading('3.2 alertas-websocket - Comunicacion en Tiempo Real', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Servidor WebSocket dedicado que mantiene conexiones persistentes con los clientes para '
    'notificaciones instantaneas. Permite que las alertas lleguen a los cuidadores en tiempo real '
    'sin necesidad de hacer polling al servidor.'
)
doc.add_paragraph('URL: https://alertas-websocket-687053793381.southamerica-west1.run.app')
doc.add_paragraph('Funcionalidades:')
bullets = [
    'Conexiones WebSocket persistentes con la aplicacion movil',
    'Notificaciones instantaneas de alertas de caidas',
    'Notificaciones de solicitudes de ayuda',
    'Actualizacion de estado de alertas en tiempo real',
    'Reconexion automatica ante perdida de conexion',
    'Soporte para multiples dispositivos por usuario'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.3 api-email
doc.add_heading('3.3 api-email (SendGrid) - Notificaciones por Correo', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Microservicio dedicado al envio de correos electronicos transaccionales. Utiliza la API de '
    'SendGrid para garantizar alta entregabilidad y seguimiento de emails.'
)
doc.add_paragraph('URL: https://api-email-687053793381.southamerica-west1.run.app')
doc.add_paragraph('Funcionalidades:')
bullets = [
    'Envio de alertas de caidas por email a todos los cuidadores configurados',
    'Notificaciones de solicitudes de ayuda',
    'Confirmaciones de registro de nuevos usuarios',
    'Plantillas HTML personalizadas con la marca VigilIA',
    'Seguimiento de entrega y apertura de emails',
    'Cola de reintentos para emails fallidos'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.4 whatsapp-webhook
doc.add_heading('3.4 whatsapp-webhook (Meta API) - Mensajeria Instantanea', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Servicio de integracion con WhatsApp Business API de Meta. Permite enviar mensajes de alerta '
    'criticos directamente al WhatsApp de los cuidadores, garantizando que reciban la notificacion '
    'incluso si no tienen la aplicacion abierta.'
)
doc.add_paragraph('URL: https://whatsapp-webhook-687053793381.southamerica-west1.run.app')
doc.add_paragraph('Funcionalidades:')
bullets = [
    'Envio de alertas criticas de caidas via WhatsApp',
    'Mensajes con plantillas aprobadas por Meta',
    'Inclusion de imagen de evidencia en el mensaje',
    'Recepcion de respuestas de confirmacion',
    'Webhook para verificacion de entrega',
    'Soporte para numeros internacionales'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 4. CLOUD SQL - BASE DE DATOS
# ============================================================
doc.add_heading('4. Cloud SQL (PostgreSQL) - Base de Datos', level=2)

doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Sistema gestor de base de datos relacional responsable de almacenar toda la informacion '
    'del sistema. PostgreSQL fue elegido por su robustez, soporte de JSON, y excelente rendimiento '
    'para consultas complejas.'
)

doc.add_paragraph('Especificaciones tecnicas de la instancia:')
bullets = [
    'Nombre de instancia: vigilia-db-main',
    'Motor: PostgreSQL 17 (ultima version estable)',
    'Tier: db-custom-2-8192 (2 vCPU, 8 GB RAM)',
    'Almacenamiento: 10 GB SSD con autoescalado',
    'Region: southamerica-west1 (Santiago, Chile)',
    'Alta disponibilidad: Backups automaticos diarios',
    'Conexion: IP privada + Cloud SQL Proxy para seguridad'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Esquema de base de datos (10 tablas en produccion):')

# Tabla de tablas
tables_info = [
    ('usuarios', 'Cuentas de usuario del sistema. Almacena cuidadores, administradores y adultos mayores con sus credenciales, roles y tokens de notificacion push.'),
    ('adultos_mayores', 'Perfiles de las personas monitoreadas. Incluye nombre, fecha de nacimiento, direccion y configuraciones especificas.'),
    ('cuidadores_adultos_mayores', 'Tabla de relacion N:N que vincula cuidadores con adultos mayores. Un cuidador puede monitorear multiples adultos y un adulto puede tener multiples cuidadores.'),
    ('dispositivos', 'Registro de dispositivos NanoPi. Almacena el ID unico, nombre, estado de conexion y adulto mayor asociado.'),
    ('configuraciones_alerta', 'Preferencias de notificacion por usuario. Define que canales usar (email, WhatsApp, push) y horarios de no molestar.'),
    ('alertas', 'Eventos de caida (tipo=caida) y solicitudes de ayuda (tipo=ayuda). Registra fecha, hora, imagen de evidencia, estado y adulto mayor involucrado.'),
    ('alertas_vistas', 'Control de lectura de alertas por usuario. Permite saber cuales alertas han sido vistas por cada cuidador.'),
    ('recordatorios', 'Recordatorios configurados: medicamentos, citas medicas, ejercicio, hidratacion. Incluye frecuencia, hora y estado activo/inactivo.'),
    ('suscripciones', 'Planes de pago contratados. Define el tipo de plan (basico, plus, premium), fecha de inicio/fin y estado de pago.'),
    ('solicitudes_cuidado', 'Invitaciones de cuidado entre usuarios. Permite que un usuario invite a otro a ser cuidador de un adulto mayor.')
]

for table_name, description in tables_info:
    p = doc.add_paragraph()
    run = p.add_run(f'{table_name}: ')
    run.bold = True
    p.add_run(description)

# ============================================================
# 5. CLOUD STORAGE
# ============================================================
doc.add_heading('5. Cloud Storage (Buckets) - Almacenamiento de Archivos', level=2)

doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Servicio de almacenamiento de objetos utilizado para guardar las evidencias visuales '
    'capturadas durante la deteccion de caidas. Proporciona almacenamiento duradero, seguro '
    'y accesible desde cualquier servicio.'
)

doc.add_paragraph('Configuracion del bucket principal:')
bullets = [
    'Nombre: nanopi-videos-input',
    'Region: southamerica-west1 (misma region que los servicios)',
    'Clase de almacenamiento: Standard (acceso frecuente)',
    'Control de acceso: Uniforme con IAM',
    'URLs firmadas para acceso temporal seguro desde la app'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Contenido almacenado:')
bullets = [
    'Snapshots (fotografias) capturadas en el momento de deteccion de caida',
    'Imagenes de evidencia asociadas a cada alerta',
    'Formato: JPEG optimizado para reducir tamano',
    'Nomenclatura: {dispositivo_id}/{timestamp}.jpg'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 6. SECRET MANAGER
# ============================================================
doc.add_heading('6. Secret Manager - Gestion de Credenciales', level=2)

doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Servicio de gestion segura de credenciales y secretos. Todas las claves API, contrasenas '
    'y tokens sensibles se almacenan aqui en lugar de en el codigo fuente, siguiendo las '
    'mejores practicas de seguridad.'
)

doc.add_paragraph('Secretos configurados y en uso:')
bullets = [
    'vigilia-db-password: Contrasena de conexion a la base de datos PostgreSQL',
    'internal-api-key: API key para comunicacion segura entre microservicios internos',
    'SENDGRID_API_KEY: Credenciales de autenticacion con SendGrid para envio de emails',
    'whatsapp-api-key: Clave de API de Meta para WhatsApp Business',
    'whatsapp-token: Token de acceso para envio de mensajes WhatsApp',
    'webhook-api-key: Clave para autenticar webhooks entrantes',
    'webhook-verify-token: Token de verificacion para validar origen de webhooks'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 7. SERVICIOS EXTERNOS
# ============================================================
doc.add_heading('7. Servicios de Notificacion (Externos)', level=2)

doc.add_paragraph(
    'El sistema integra tres canales de notificacion externos para garantizar que las alertas '
    'lleguen a los cuidadores por multiples vias, aumentando la probabilidad de respuesta rapida.'
)

doc.add_heading('7.1 SendGrid - Email Transaccional', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
bullets = [
    'Proveedor: Twilio SendGrid',
    'Uso: Notificaciones de alertas y confirmaciones por correo electronico',
    'Entregabilidad: Mayor al 95% en bandeja de entrada',
    'Plantillas HTML responsivas para lectura en movil y desktop'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('7.2 Meta WhatsApp Business API', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
bullets = [
    'Proveedor: Meta (Facebook)',
    'Uso: Alertas criticas de caidas via WhatsApp',
    'Ventaja: Notificacion inmediata incluso con app cerrada',
    'Incluye imagen de evidencia en el mensaje'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('7.3 Expo Push Notifications', level=3)
doc.add_paragraph('Estado actual: OPERATIVO')
bullets = [
    'Proveedor: Expo (para React Native)',
    'Uso: Notificaciones push nativas en Android e iOS',
    'Alertas instantaneas de caidas y solicitudes de ayuda',
    'Recordatorios programados de medicamentos y citas',
    'Funciona incluso con la aplicacion en segundo plano'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 8. DISPOSITIVOS EDGE (ON-PREMISE)
# ============================================================
doc.add_heading('8. Dispositivos Edge (On-Premise)', level=2)

doc.add_paragraph(
    'Componentes de hardware instalados en el hogar del adulto mayor para la deteccion local '
    'de caidas. El procesamiento se realiza localmente, enviando solo alertas a la nube, '
    'lo que protege la privacidad y reduce el consumo de ancho de banda.'
)

doc.add_heading('8.1 NanoPi Neo4 - Computadora de Borde', level=3)
doc.add_paragraph('Estado actual: OPERATIVO (dispositivos en pruebas de campo)')
doc.add_paragraph(
    'Dispositivo de computacion edge (Single-Board Computer) que ejecuta el analisis de video '
    'y la deteccion de caidas de forma autonoma, sin depender de la nube para el procesamiento.'
)
doc.add_paragraph('Especificaciones tecnicas:')
bullets = [
    'Procesador: ARM64 Rockchip RK3399 (6 nucleos)',
    'RAM: 4GB DDR3',
    'Sistema operativo: Ubuntu Server 22.04 LTS',
    'Consumo energetico: Menor a 10W',
    'Conectividad: Ethernet Gigabit + WiFi opcional'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Software instalado y funcionando:')
bullets = [
    'Python 3.10 con entorno virtual dedicado',
    'MediaPipe: Biblioteca de Google para deteccion de poses corporales',
    'OpenCV: Procesamiento de video y captura de frames',
    'Modelo de deteccion de caidas entrenado (.pt)',
    'Cliente HTTPS para comunicacion segura con api-backend',
    'Script de inicio automatico al encender el dispositivo'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Flujo de deteccion:')
bullets = [
    '1. Recibe stream RTSP de la camara IP',
    '2. Procesa cada frame con MediaPipe para detectar poses',
    '3. Analiza posicion del cuerpo con el modelo de IA',
    '4. Si detecta caida: captura snapshot y envia alerta a api-backend',
    '5. El video NUNCA sale del hogar, solo la alerta y la imagen'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('8.2 Camara IP Dahua - Captura de Video', level=3)
doc.add_paragraph('Estado actual: COMPATIBLE Y PROBADA')
doc.add_paragraph(
    'Camara de vigilancia IP que proporciona el stream de video al NanoPi. Se comunica '
    'exclusivamente dentro de la red local del hogar.'
)
doc.add_paragraph('Especificaciones:')
bullets = [
    'Marca recomendada: Dahua (modelos IPC-HDW series)',
    'Protocolo: RTSP para transmision de video',
    'Resolucion: Substream 640x480 (optimizado para deteccion)',
    'Conexion: Red local via cable Ethernet o WiFi',
    'Sin acceso desde internet (seguridad por diseno)'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 9. APLICACION MOVIL Y WEB
# ============================================================
doc.add_heading('9. Aplicacion Movil y Web (Frontend)', level=2)

doc.add_paragraph('Estado actual: OPERATIVO')
doc.add_paragraph(
    'Interfaz de usuario desarrollada con React Native y Expo, permitiendo una unica base '
    'de codigo para web, Android e iOS. La aplicacion esta disponible tanto en navegador '
    'como en dispositivos moviles.'
)

doc.add_paragraph('Tecnologias utilizadas:')
bullets = [
    'Framework: React Native con Expo SDK 52',
    'Lenguaje: TypeScript para tipado estatico',
    'Navegacion: Expo Router (file-based routing)',
    'Estado: React Context API',
    'Estilos: StyleSheet nativo de React Native'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Plataformas disponibles:')
bullets = [
    'Web: https://app-vigilia.web.app (Firebase Hosting)',
    'Android: APK generado via EAS Build',
    'iOS: Build disponible via EAS Build (requiere cuenta Apple Developer)'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Funcionalidades implementadas en la aplicacion:')
bullets = [
    'Inicio de sesion y registro de usuarios',
    'Dashboard con resumen de alertas y recordatorios',
    'Visualizacion de alertas de caidas con imagen de evidencia',
    'Visualizacion de solicitudes de ayuda',
    'Gestion de recordatorios (crear, editar, eliminar)',
    'Perfil de usuario y configuracion de notificaciones',
    'Seleccion de adulto mayor a monitorear',
    'Invitacion a otros cuidadores',
    'Graficos de actividad y estadisticas',
    'Notificaciones push en tiempo real'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# IMAGEN PLACEHOLDER
# ============================================================
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[INSERTAR DIAGRAMA DE DESPLIEGUE AQUI]')
run.italic = True

doc.add_paragraph('')
p = doc.add_paragraph('Imagen 7. Diagrama de despliegue, muestra componentes del sistema. Fuente: mivigilia.cl')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# RESUMEN DE URLS Y SERVICIOS
# ============================================================
doc.add_heading('Resumen de URLs y Servicios en Produccion', level=2)

doc.add_paragraph('Todos los siguientes servicios estan actualmente operativos:')

services = [
    ('Frontend Web', 'https://app-vigilia.web.app', 'Firebase Hosting'),
    ('API Principal', 'https://api-backend-687053793381.southamerica-west1.run.app', 'Cloud Run'),
    ('WebSocket', 'https://alertas-websocket-687053793381.southamerica-west1.run.app', 'Cloud Run'),
    ('API Email', 'https://api-email-687053793381.southamerica-west1.run.app', 'Cloud Run'),
    ('WhatsApp Webhook', 'https://whatsapp-webhook-687053793381.southamerica-west1.run.app', 'Cloud Run'),
    ('Base de Datos', 'vigilia-db-main (IP privada)', 'Cloud SQL'),
    ('Almacenamiento', 'gs://nanopi-videos-input', 'Cloud Storage')
]

for name, url, platform in services:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(f'{url} ({platform})')

# Guardar documento
output_path = r'C:\Users\acuri\Documents\Vigilia\Capstone\docs\seccion-diagrama-despliegue-COMPLETA.docx'
doc.save(output_path)
print(f'Documento creado exitosamente en: {output_path}')
