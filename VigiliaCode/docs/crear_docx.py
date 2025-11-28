# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titulo
title = doc.add_heading('Diagrama de Despliegue', level=1)

# Intro
doc.add_paragraph('Este diagrama muestra como se distribuye fisicamente el sistema en su infraestructura, donde se ejecuta cada componente. Ayuda y facilita la comunicacion entre los distintos equipos de desarrollo, ademas de analizar el rendimiento, la disponibilidad y la escalabilidad del sistema.')

doc.add_paragraph('En el diagrama se muestran los principales componentes que permiten la ejecucion optima e integrada de todo el sistema:')

# 1. GCP
doc.add_heading('Google Cloud Platform (GCP)', level=2)
doc.add_paragraph('La plataforma en la cual se aloja toda la infraestructura del sistema. Permite un despliegue seguro, escalable y con servicios integrados que simplifican la operacion y el mantenimiento. GCP provee el ecosistema donde se ejecutan los servicios, bases de datos, almacenamiento y mensajeria.')
doc.add_paragraph('Region principal: southamerica-west1 (Santiago, Chile)')
doc.add_paragraph('Proyecto: composed-apogee-475623-p6')

# 2. Firebase
doc.add_heading('Firebase Services', level=2)

doc.add_heading('Firebase Hosting', level=3)
doc.add_paragraph('Servicio de alojamiento web donde se despliega el frontend de la aplicacion. Proporciona:')
bullets = ['HTTPS automatico con certificado SSL', 'CDN global para baja latencia', 'Despliegue continuo desde el repositorio', 'URL: app-vigilia.web.app']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Firebase Authentication', level=3)
doc.add_paragraph('Sistema de autenticacion de usuarios que maneja:')
bullets = ['Registro e inicio de sesion con email/password', 'Generacion y validacion de tokens JWT', 'Integracion nativa con el frontend React Native']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3. Cloud Run
doc.add_heading('Cloud Run (Servicios Serverless)', level=2)
doc.add_paragraph('Plataforma serverless que ejecuta los microservicios del backend en contenedores Docker con escalamiento automatico. Caracteristicas principales:')
bullets = ['Escalamiento automatico de 0 a N instancias segun demanda', 'Pago solo por uso (sin costo cuando no hay trafico)', 'Aislamiento seguro mediante contenedores', 'Despliegue continuo desde Cloud Build']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Microservicios desplegados:')

# api-backend
doc.add_heading('api-backend (FastAPI)', level=3)
doc.add_paragraph('Componente central del procesamiento logico del sistema. Sus funciones principales son:')
bullets = ['Exponer endpoints REST para comunicacion con el frontend', 'Procesar informacion recibida desde los dispositivos NanoPi', 'Ejecutar logica de negocio, validaciones y generacion de respuestas', 'Integracion con la base de datos, almacenamiento y servicios de mensajeria', 'Gestion de usuarios, alertas, recordatorios y dispositivos']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# alertas-websocket
doc.add_heading('alertas-websocket', level=3)
doc.add_paragraph('Servidor WebSocket para comunicacion en tiempo real:')
bullets = ['Notificaciones instantaneas de alertas a la aplicacion', 'Conexion persistente con los clientes', 'Actualizacion de estado en tiempo real']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# api-email
doc.add_heading('api-email (SendGrid)', level=3)
doc.add_paragraph('Microservicio dedicado al envio de correos electronicos:')
bullets = ['Notificaciones de alertas por email', 'Confirmaciones de registro', 'Integracion con SendGrid API']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# whatsapp-webhook
doc.add_heading('whatsapp-webhook (Meta API)', level=3)
doc.add_paragraph('Servicio de integracion con WhatsApp Business:')
bullets = ['Envio de alertas criticas por WhatsApp', 'Recepcion de mensajes de confirmacion', 'Integracion con Meta WhatsApp Cloud API']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 4. Cloud SQL
doc.add_heading('Cloud SQL (PostgreSQL)', level=2)
doc.add_paragraph('Sistema gestor de base de datos relacional responsable de almacenar toda la informacion del sistema.')
doc.add_paragraph('Especificaciones tecnicas:')
bullets = ['Instancia: vigilia-db-main', 'Version: PostgreSQL 17', 'Tier: db-custom-2-8192 (2 vCPU, 8GB RAM)', 'Region: southamerica-west1']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Tablas de la base de datos (10 tablas):')
bullets = [
    'usuarios - Cuentas de usuario (cuidadores, administradores, adultos mayores)',
    'adultos_mayores - Perfiles de personas monitoreadas',
    'cuidadores_adultos_mayores - Relacion N:N entre cuidadores y adultos mayores',
    'dispositivos - Registro de dispositivos NanoPi y camaras',
    'configuraciones_alerta - Preferencias de notificacion por usuario',
    'alertas - Eventos de caida (tipo=caida) y solicitudes de ayuda (tipo=ayuda)',
    'alertas_vistas - Estado de lectura de alertas/recordatorios por usuario',
    'recordatorios - Medicamentos, citas medicas, ejercicio, etc.',
    'suscripciones - Planes de pago (basico, plus, premium)',
    'solicitudes_cuidado - Invitaciones de cuidado entre usuarios'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 5. Cloud Storage
doc.add_heading('Cloud Storage (Buckets)', level=2)
doc.add_paragraph('Servicio de almacenamiento de objetos utilizado para guardar:')
bullets = ['Fotografias (snapshots) capturadas durante la deteccion de caidas', 'Archivos de evidencia asociados a alertas', 'Bucket principal: nanopi-videos-input']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 6. Secret Manager
doc.add_heading('Secret Manager', level=2)
doc.add_paragraph('Servicio de gestion segura de credenciales y secretos:')
bullets = [
    'vigilia-db-password - Contrasena de la base de datos',
    'internal-api-key - API key para comunicacion entre servicios',
    'SENDGRID_API_KEY - Credenciales de SendGrid',
    'whatsapp-api-key, whatsapp-token - Credenciales de Meta WhatsApp',
    'webhook-api-key, webhook-verify-token - Tokens de verificacion'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 7. Servicios Externos
doc.add_heading('Servicios de Mensajeria (Externos)', level=2)
doc.add_paragraph('Encargados de enviar notificaciones y alertas en tiempo real hacia los cuidadores.')

doc.add_heading('SendGrid', level=3)
doc.add_paragraph('Servicio de email transaccional para:')
bullets = ['Notificaciones de alertas por correo electronico', 'Confirmacion de cuentas y acciones']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Meta WhatsApp Business API', level=3)
doc.add_paragraph('Plataforma de mensajeria para:')
bullets = ['Enviar alertas criticas de caidas via WhatsApp', 'Confirmacion de recepcion de alertas', 'Comunicacion inmediata con cuidadores']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Expo Push Notifications', level=3)
doc.add_paragraph('Sistema de notificaciones push para la aplicacion movil:')
bullets = ['Alertas instantaneas en dispositivos Android e iOS', 'Recordatorios programados', 'Actualizaciones de estado']
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 8. Dispositivos Edge
doc.add_heading('Dispositivos Edge (On-Premise)', level=2)
doc.add_paragraph('Componentes instalados en el hogar del adulto mayor para la deteccion local de caidas.')

doc.add_heading('NanoPi Neo4', level=3)
doc.add_paragraph('Dispositivo de computacion edge (Single-Board Computer) que ejecuta el analisis de video:')
bullets = [
    'Procesador ARM64 de bajo consumo',
    'Sistema operativo Ubuntu',
    'Software de deteccion con MediaPipe y Python',
    'Conexion a la nube via HTTPS',
    'Analisis de video en tiempo real sin enviar streaming'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Camara IP Dahua', level=3)
doc.add_paragraph('Camara de vigilancia que proporciona el stream de video:')
bullets = [
    'Protocolo RTSP para transmision local',
    'Resolucion configurable (substream para optimizar recursos)',
    'Conexion directa con el NanoPi via red local'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Imagen placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[INSERTAR DIAGRAMA DE DESPLIEGUE AQUI]')
run.italic = True

doc.add_paragraph('')
p = doc.add_paragraph('Imagen 7. Diagrama de despliegue, muestra componentes del sistema. Fuente: mivigilia.cl')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Guardar
doc.save(r'C:\Users\acuri\Documents\Vigilia\Capstone\docs\seccion-diagrama-despliegue-actualizado.docx')
print('Documento creado exitosamente!')
